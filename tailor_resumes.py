#!/usr/bin/env python3
"""Generate tailored resume variants from base resume using python-docx.

Each profile targets a cluster of Tier A jobs with similar ATS keywords.
Only REFRAMES existing text — never adds paragraphs or sections.

SETUP INSTRUCTIONS:
1. Set BASE_DOCX to the path of your base resume .docx file
2. Set OUTPUT_DIR to today's date folder (updated by Claude Code each run)
3. Define your PROFILES below — each profile is a set of paragraph-index
   replacements that reframe your resume for a specific job cluster

To find your paragraph indices, run:
    python -c "from docx import Document; d=Document('your_resume.docx'); [print(f'{i}: {p.text[:80]}') for i,p in enumerate(d.paragraphs)]"
"""

import json
import os
from datetime import date
from docx import Document


# ── Configuration ─────────────────────────────────────────────────────────────
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
try:
    with open(os.path.join(SCRIPT_DIR, "config.json"), encoding="utf-8") as _cf:
        _config = json.load(_cf)
    BASE_DOCX = _config["resume"]["base_docx_path"]
except (FileNotFoundError, KeyError, json.JSONDecodeError) as e:
    print(f"ERROR: config.json missing or malformed: {e}")
    print("Create config.json with resume.base_docx_path set to your base resume path.")
    BASE_DOCX = None
OUTPUT_DIR = os.path.join(SCRIPT_DIR, "tailored-resumes", date.today().isoformat())


# ── Helpers ──────────────────────────────────────────────────────────────────

def replace_paragraph_text(para, new_text):
    """Replace all text in paragraph, preserving first run's formatting."""
    if not para.runs:
        return
    para.runs[0].text = new_text
    for run in para.runs[1:]:
        run.text = ""


def replace_label_value(para, new_label=None, new_value=None):
    """Replace bold label and/or non-bold value in a label:value paragraph.

    Works for lines like:  **Skills:** Python, SQL, Tableau
    Where "Skills:" is bold and the rest is not.
    """
    bold_runs = [r for r in para.runs if r.font.bold]
    value_runs = [r for r in para.runs if not r.font.bold]

    if new_label and bold_runs:
        bold_runs[0].text = new_label
        for r in bold_runs[1:]:
            r.text = ""

    if new_value and value_runs:
        value_runs[0].text = new_value
        for r in value_runs[1:]:
            r.text = ""
    elif new_value and not value_runs and bold_runs:
        # All runs are bold (e.g., project titles) — replace via last bold run
        bold_runs[-1].text = new_value
        for r in bold_runs[:-1]:
            r.text = ""


def apply_profile(doc, profile):
    """Apply a profile's paragraph replacements to a Document."""
    paras = doc.paragraphs

    for idx, change in profile["changes"].items():
        p = paras[idx]
        if isinstance(change, str):
            replace_paragraph_text(p, change)
        elif isinstance(change, dict):
            replace_label_value(p, change.get("label"), change.get("value"))


# ── Profile Definitions ─────────────────────────────────────────────────────
# Each profile has:
#   name: filename stem (e.g., "Product_Analytics" -> Product_Analytics.docx)
#   changes: {paragraph_index: new_text_or_label_value_dict}
#
# EXAMPLE profiles for a Data Analyst — customize these for YOUR resume:
#
# To find your paragraph indices:
#   python -c "from docx import Document; d=Document('resume.docx'); [print(f'{i}: {p.text[:80]}') for i,p in enumerate(d.paragraphs)]"
#
# Then define profiles like:
#
# PROFILES = [
#     {
#         "name": "Product_Analytics",
#         "changes": {
#             # Summary paragraph (update index to match your resume)
#             3: "Product-focused Data Analyst with expertise in A/B testing, SQL, and Tableau...",
#             # Skills line — reorder to emphasize product analytics tools
#             7: {"label": "Tools: ", "value": "SQL, Python, Tableau, Amplitude, Mixpanel, Excel"},
#             # Work experience bullet — reframe with product analytics language
#             12: "Designed and maintained product analytics dashboards tracking user engagement...",
#         },
#     },
#     {
#         "name": "BI_Reporting",
#         "changes": {
#             3: "Business Intelligence Analyst specializing in dashboard development and reporting...",
#             7: {"label": "Tools: ", "value": "Power BI, Tableau, SQL, Excel, Looker, dbt"},
#             12: "Built automated BI dashboards reducing manual reporting time by 60%...",
#         },
#     },
# ]

PROFILES = []  # ← Define your profiles here (see examples above)


# ── Generate ─────────────────────────────────────────────────────────────────

def main():
    if not BASE_DOCX:
        return 1
    if not os.path.exists(BASE_DOCX):
        print(f"ERROR: Base resume not found: {BASE_DOCX}")
        return 1

    if not PROFILES:
        print("No profiles defined in tailor_resumes.py.")
        print("See the comments in the file for setup instructions.")
        print("\nTo find your paragraph indices, run:")
        print(f'  python -c "from docx import Document; d=Document(\'{BASE_DOCX}\'); [print(f\'{{i}}: {{p.text[:80]}}\') for i,p in enumerate(d.paragraphs)]"')
        return

    os.makedirs(OUTPUT_DIR, exist_ok=True)

    for profile in PROFILES:
        doc = Document(BASE_DOCX)
        apply_profile(doc, profile)
        out_path = os.path.join(OUTPUT_DIR, f"{profile['name']}.docx")
        doc.save(out_path)
        print(f"  Created: {profile['name']}.docx")

    print(f"\nAll {len(PROFILES)} tailored resumes saved to {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
